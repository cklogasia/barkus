#!/usr/bin/env python3
"""
Script to create a DOCX document with Windows 11 compilation guide for Barkus
"""

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_compilation_guide():
    # Create a new document
    doc = Document()
    
    # Add title
    title = doc.add_heading('Windows 11 Compilation Guide for Barkus', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add subtitle
    subtitle = doc.add_paragraph('PDF Barcode Splitter - Complete Setup for Fresh Environment')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_format = subtitle.runs[0]
    subtitle_format.italic = True
    
    doc.add_paragraph()  # Add spacing
    
    # Overview section
    doc.add_heading('Overview', level=1)
    overview_text = """Barkus is a Python application that processes PDF documents to detect delivery number and customer name barcodes on each page, then splits the PDF into separate files based on these barcode combinations. This guide provides step-by-step instructions for compiling the application on a fresh Windows 11 environment with no existing Python installation."""
    doc.add_paragraph(overview_text)
    
    # Prerequisites section
    doc.add_heading('System Requirements', level=1)
    req_list = doc.add_paragraph()
    req_list.add_run('• Windows 11 (64-bit recommended)\n')
    req_list.add_run('• Administrator privileges for installation\n')
    req_list.add_run('• Internet connection for downloading dependencies\n')
    req_list.add_run('• At least 2GB free disk space')
    
    # Step 1: Python Installation
    doc.add_heading('Step 1: Install Python 3.7+', level=1)
    
    doc.add_heading('Download and Install Python', level=2)
    python_steps = doc.add_paragraph()
    python_steps.add_run('1. Visit ').bold = False
    python_steps.add_run('https://python.org/downloads/').bold = True
    python_steps.add_run('\n2. Download Python 3.11 (recommended) for Windows\n')
    python_steps.add_run('3. Run the installer with these ')
    python_steps.add_run('IMPORTANT').bold = True
    python_steps.add_run(' options:\n')
    python_steps.add_run('   ✓ Add Python to PATH\n')
    python_steps.add_run('   ✓ Install pip\n')
    python_steps.add_run('   ✓ Install for all users (optional)\n')
    python_steps.add_run('4. Verify installation by opening Command Prompt and running:')
    
    code_block = doc.add_paragraph('python --version\npip --version')
    code_block.style = 'Intense Quote'
    
    # Step 2: System Dependencies
    doc.add_heading('Step 2: Install System Dependencies', level=1)
    
    doc.add_heading('Install Poppler (CRITICAL)', level=2)
    poppler_text = doc.add_paragraph()
    poppler_text.add_run('Poppler is ').bold = False
    poppler_text.add_run('REQUIRED').bold = True
    poppler_text.add_run(' for pdf2image to work:\n\n')
    poppler_text.add_run('1. Download poppler-windows from:\n   ')
    poppler_text.add_run('https://github.com/oschwartz10612/poppler-windows/releases/').bold = True
    poppler_text.add_run('\n2. Extract the ZIP file to ')
    poppler_text.add_run('C:\\poppler').bold = True
    poppler_text.add_run('\n3. Add ')
    poppler_text.add_run('C:\\poppler\\Library\\bin').bold = True
    poppler_text.add_run(' to system PATH:\n')
    poppler_text.add_run('   • Right-click "This PC" → Properties\n')
    poppler_text.add_run('   • Advanced system settings\n')
    poppler_text.add_run('   • Environment Variables\n')
    poppler_text.add_run('   • Edit "Path" in System variables\n')
    poppler_text.add_run('   • Add new entry: C:\\poppler\\Library\\bin')
    
    doc.add_heading('Install Visual C++ Build Tools', level=2)
    vc_text = doc.add_paragraph()
    vc_text.add_run('Required for compiling some Python packages:\n\n')
    vc_text.add_run('Option 1: ')
    vc_text.add_run('Microsoft C++ Build Tools').bold = True
    vc_text.add_run('\n• Download from Microsoft Developer site\n')
    vc_text.add_run('• Install with C++ build tools workload\n\n')
    vc_text.add_run('Option 2: ')
    vc_text.add_run('Visual Studio Community').bold = True
    vc_text.add_run('\n• Install with "Desktop development with C++" workload')
    
    # Step 3: Project Setup
    doc.add_heading('Step 3: Project Setup and Dependencies', level=1)
    
    doc.add_heading('Create Project Environment', level=2)
    setup_commands = doc.add_paragraph('Navigate to your project directory and run:')
    
    code_setup = doc.add_paragraph('''# Navigate to project directory
cd path\\to\\your\\barkus\\project

# Create virtual environment
python -m venv venv

# Activate virtual environment
venv\\Scripts\\activate

# Upgrade pip
pip install --upgrade pip

# Install project dependencies
pip install -r requirements.txt''')
    code_setup.style = 'Intense Quote'
    
    # Dependencies breakdown
    doc.add_heading('Key Dependencies Explained', level=2)
    deps_table_data = [
        ['Package', 'Purpose', 'Special Requirements'],
        ['pikepdf', 'PDF manipulation and processing', 'None'],
        ['pdf2image', 'Convert PDF pages to images', 'Requires Poppler'],
        ['opencv-python', 'Computer vision for barcode detection', 'None'],
        ['pyzbar', 'Barcode reading and decoding', 'May need libzbar DLL'],
        ['numpy', 'Numerical operations', 'Version <2.0.0 required']
    ]
    
    table = doc.add_table(rows=len(deps_table_data), cols=3)
    table.style = 'Table Grid'
    
    for i, row_data in enumerate(deps_table_data):
        row = table.rows[i]
        for j, cell_data in enumerate(row_data):
            row.cells[j].text = cell_data
            if i == 0:  # Header row
                row.cells[j].paragraphs[0].runs[0].bold = True
    
    # Step 4: Compilation
    doc.add_heading('Step 4: Compile to Executable', level=1)
    
    doc.add_heading('Install PyInstaller', level=2)
    pyinst_text = doc.add_paragraph('Install the PyInstaller package:')
    
    pyinst_code = doc.add_paragraph('pip install pyinstaller')
    pyinst_code.style = 'Intense Quote'
    
    doc.add_heading('Compilation Options', level=2)
    
    doc.add_heading('Option 1: Single File Executable', level=3)
    single_file_text = doc.add_paragraph('Creates a single .exe file (slower startup):')
    single_file_code = doc.add_paragraph('pyinstaller --onefile --add-data "poppler;poppler" barkus.py')
    single_file_code.style = 'Intense Quote'
    
    doc.add_heading('Option 2: Directory Distribution (Recommended)', level=3)
    dir_dist_text = doc.add_paragraph('Creates a folder with executable and dependencies (faster startup):')
    dir_dist_code = doc.add_paragraph('pyinstaller --onedir --add-data "C:\\poppler\\Library\\bin;poppler\\bin" barkus.py')
    dir_dist_code.style = 'Intense Quote'
    
    # Step 5: Testing
    doc.add_heading('Step 5: Testing and Distribution', level=1)
    
    doc.add_heading('Test the Executable', level=2)
    test_steps = doc.add_paragraph()
    test_steps.add_run('1. Navigate to ')
    test_steps.add_run('dist\\barkus\\').bold = True
    test_steps.add_run(' directory\n')
    test_steps.add_run('2. Run ')
    test_steps.add_run('barkus.exe --help').bold = True
    test_steps.add_run(' to verify it works\n')
    test_steps.add_run('3. Test with a sample PDF file\n')
    test_steps.add_run('4. Verify barcode detection and PDF splitting functionality')
    
    doc.add_heading('Distribution', level=2)
    dist_text = doc.add_paragraph()
    dist_text.add_run('The compiled application can now run on any Windows 11 machine ')
    dist_text.add_run('without requiring Python installation').bold = True
    dist_text.add_run('. Simply copy the entire ')
    dist_text.add_run('dist\\barkus\\').bold = True
    dist_text.add_run(' folder to the target machine.')
    
    # Troubleshooting section
    doc.add_heading('Troubleshooting', level=1)
    
    troubleshoot_items = [
        ("'poppler' not found error", "Ensure Poppler is installed and C:\\poppler\\Library\\bin is in your system PATH. Restart Command Prompt after adding to PATH."),
        ("Import errors during compilation", "Make sure all dependencies are installed in your virtual environment. Run 'pip list' to verify."),
        ("Missing DLL errors", "Install Visual C++ Redistributable packages on the target machine if needed."),
        ("Barcode detection not working", "Verify that the input PDF contains valid barcodes and try adjusting the DPI parameter (--dpi flag).")
    ]
    
    for issue, solution in troubleshoot_items:
        doc.add_heading(issue, level=2)
        doc.add_paragraph(solution)
    
    # Usage section
    doc.add_heading('Usage Examples', level=1)
    
    usage_text = doc.add_paragraph('Once compiled, use the executable with these commands:')
    
    usage_examples = doc.add_paragraph('''# Basic usage
barkus.exe input.pdf

# Specify output directory
barkus.exe input.pdf --output-dir C:\\output

# Handle pages without barcodes
barkus.exe input.pdf --handle-no-barcode separate

# Adjust DPI for better barcode detection
barkus.exe input.pdf --dpi 600

# Quiet mode
barkus.exe input.pdf --quiet''')
    usage_examples.style = 'Intense Quote'
    
    # Save the document
    doc.save('Barkus_Windows11_Compilation_Guide.docx')
    print("DOCX guide created successfully: Barkus_Windows11_Compilation_Guide.docx")

if __name__ == "__main__":
    create_compilation_guide()